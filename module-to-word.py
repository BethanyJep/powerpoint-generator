import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin
import time

headers = {"User-Agent": "Mozilla/5.0"}
base_url = "https://learn.microsoft.com/en-us/training/modules/develop-ai-agent-with-semantic-kernel/"

def get_soup(url):
    response = requests.get(url, headers=headers)
    return BeautifulSoup(response.content, "html.parser")

def extract_module_unit_links(start_url):
    soup = get_soup(start_url)
    nav_section = soup.find("ul", {"id": "unit-list"})

    if not nav_section:
        print("❌ Could not find unit navigation.")
        return []

    links = []
    for a in nav_section.find_all("a", href=True):
        href = a["href"]
        full_url = urljoin(base_url, href)
        links.append(full_url)

    return list(dict.fromkeys(links))  # remove duplicates while preserving order

def extract_content(soup, doc):
    main_content = soup.find("main")
    if not main_content:
        return

    for tag in main_content.find_all(["h1", "h2", "h3", "p", "li", "pre", "code"]):
        text = tag.get_text().strip()
        if not text:
            continue
        if tag.name == "h1":
            doc.add_heading(text, level=1)
        elif tag.name == "h2":
            doc.add_heading(text, level=2)
        elif tag.name == "h3":
            doc.add_heading(text, level=3)
        elif tag.name == "p":
            doc.add_paragraph(text)
        elif tag.name == "li":
            doc.add_paragraph(f"• {text}", style='ListBullet')
        elif tag.name in ["pre", "code"]:
            doc.add_paragraph(text, style='Intense Quote')

def scrape_full_module(start_url, output_filename="Learn_Module.docx"):
    doc = Document()
    all_unit_links = extract_module_unit_links(start_url)

    if not all_unit_links:
        print("❌ No unit links found. Exiting.")
        return

    print(f"🔗 Found {len(all_unit_links)} unit pages.")

    for i, url in enumerate(all_unit_links, start=1):
        print(f"📄 Scraping page {i}: {url}")
        soup = get_soup(url)
        extract_content(soup, doc)
        time.sleep(1)  # polite delay

    doc.save(output_filename)
    print(f"\n✅ Saved module to: {output_filename}")

# 🟡 Replace this with any Learn module start page
start_page = "https://learn.microsoft.com/en-us/training/modules/develop-ai-agent-with-semantic-kernel/"
scrape_full_module(start_page, "develop-ai-agent-with-semantic-kernel.docx")
